
import pytesseract 
import cv2
import numpy as np
from scipy import signal, ndimage
from PIL import Image
from pdf2image import convert_from_path

def load_im(path):
    return np.asarray(Image.open(path))/255.0

def save(path, img):
    tmp = np.asarray(img*255.0, dtype=np.uint8)
    Image.fromarray(tmp).save(path)

inp_path = 'test.jpg'
out_path = 'output.png'
#inp_path = 'out.png'
#inp_path = 'abc/test.pdf/0_test.pdf.jpg'
ix,iy = -1,-1
x1=0
y1=0
x0=0
y0=0
n=0
inp = load_im(inp_path)
#out = denoise_im_with_back(inp)
##mask=inp[]
##inp=cv2.imread(inp_path)
def draw_circle(event,x,y,flags,param):
    global ix,iy,x0,y0,x1,y1,n
    if event == cv2.EVENT_LBUTTONDBLCLK:
        n+=1
        print((x,y))
        if(n==1):
            x0,y0 = x,y
        else:
            x1,y1 = x,y
            

cv2.namedWindow('image')
cv2.setMouseCallback('image',draw_circle)
##cv2.imshow('image',img)

while(1):
    if(n==2):
        break
    cv2.imshow('image',inp)
    k = cv2.waitKey(20) & 0xFF
    if k == 27:
        break
    elif k == ord('a'):
        print( ix,iy)
cv2.destroyAllWindows()
nnn=0
diff=0
#diff==int(y1)-int(y0)
print(diff," ",y1, " ",y0)
from docx import Document
from docx.shared import Inches

document = Document()
pytesseract.pytesseract.tesseract_cmd = r"H:\Tesseract-OCR\tesseract.exe"

while(1):
    if((y1+(y1-y0)*nnn)>np.shape (inp)[0]):
        newimg=inp[(y0+diff*nnn):,:]
        break
    newimg=inp[(y0+(y1-y0)*nnn):(y1+(y1-y0)*nnn),:]
    nnn+=1
##cv2.imshow('image2',newimg)
    save('mask.jpg', newimg)
    str1=pytesseract.image_to_string(Image.open('mask.jpg'))
    print (str1,nnn,diff)
    document.add_paragraph(str1)
    document.save('demo.docx')
###print (pytesseract.image_to_string(Image.open('output.jpg'), lang='eng'))
newimg=inp[(y0+diff*nnn):(y1+diff*nnn),:]
##cv2.imshow('image2',newimg)
save('mask.jpg', newimg)
str1=pytesseract.image_to_string(Image.open('mask.jpg'))
print (str1)
document.add_paragraph(str1)

##document.add_heading('Document Title', 0)

document.save('demo.docx')
